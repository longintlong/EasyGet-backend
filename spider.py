#coding=utf-8
import os
import re #正则
import json #处理json数据
import requests #发起请求得到网页返回内容
from lxml import etree #lxml一种解析方式
import codecs 
import sys
import zipfile #生成压缩包
import docx #docx
reload(sys)
sys.setdefaultencoding( "utf-8" )
class BaiduWK(object):
    def __init__(self, url):
        self.title = None
        self.url = url
        self.docType = None
        self.headers = {
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_12_4) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/62.0.3202.75 Safari/537.36'}
        self.get_response_content(self.url)
        self.get_doc_type_and_title()

    def get_response_content(self, url):
        try:
            response = requests.get(url, headers=self.headers)
            return response.content
        except Exception as e:
            print(e)
            pass

    def get_doc_type_and_title(self):
        # 获取源码
        source_html = self.get_response_content(self.url)
        # 解析源码
        content = source_html.decode('gbk')
        # 获取文档类型
        self.docType = re.findall(r"docType.*?\:.*?\'(.*?)\'\,", content)[0]
        # 获取文档标题
        self.title = re.findall(r"title.*?\:.*?\'(.*?)\'\,", content)[0]


# 创建获取txt的类
class BDWKTXT(BaiduWK):
    def __init__(self, url):
        super(BDWKTXT,self).__init__(url)
        self.docId = None
        pass

    def get_txt(self, url):
        # 获取源码
        source_html = self.get_response_content(url)
        content = source_html.decode("gbk")
        # 获取docId
        self.docId = re.findall(r"docId.*?(\w{24}?)\'\,", content)[0]
        # 拼接请求url
        token_url = "https://wenku.baidu.com/api/doc/getdocinfo?callback=cb&doc_id=" + self.docId
        # 再次请求
        first_json = self.get_response_content(token_url).decode()
        str_first_json = re.match(r'.*?\((\{.*?\})\).*', first_json).group(1)
        the_first_json = json.loads(str_first_json)
        md5sum = the_first_json["md5sum"]
        rn = the_first_json["docInfo"]["totalPageNum"]
        rsign = the_first_json["rsign"]

        target_url = "https://wkretype.bdimg.com/retype/text/" + self.docId + "?" + md5sum + "&callback=cb" + "&pn=1&rn=" + rn + "&type=txt" + "&rsign=" + rsign

        sec_json = self.get_response_content(target_url).decode()

        str_sec_json = re.match(r'.*?\(\[(.*)\]\)$', sec_json).group(1)
        str_sec_json += ","
        str_json_list = str_sec_json.split('},')
        result_txt = ""
        str_json_list = str_json_list[:-1]
        for str_json in str_json_list:
            str_json += "}"
            pure_dic = json.loads(str_json)
            pure_txt = pure_dic["parags"][0]["c"].strip()
            result_txt += pure_txt


        # 创建文件目录
        try:
            path =  os.sep +'data'+os.sep+'www'+os.sep+'static'+os.sep+ self.docType
            os.makedirs(path)
            print 'erroe'
        except Exception as e:
            print 'erroe'
            pass

        # 创建文件,保存信息
        try:
            print len(result_txt)
            file_name = os.sep +'data'+os.sep+'www'+os.sep+'static'+os.sep+ self.docType+os.sep + self.title + '.docx'
            file=docx.Document()
            file.add_paragraph(result_txt)
            file.save(file_name)
            return os.sep+'static'+os.sep+ self.docType+os.sep + self.title + '.docx'
            # file_name = "." + os.sep + self.docType + os.sep + self.title + ".txt"
            # with codecs.open(file_name, 'w', 'utf-8') as f:
            #     f.write(result_txt)
            #     print("%s %s" % ("已经保存为:".decode('utf8'), self.title + '.txt'))
        except Exception as e:
            return '0'
            pass


# 创建获取word的类
class BDWKDOC(BaiduWK):
    def __init__(self, url):
        super(BDWKDOC,self).__init__(url)
        # 保存数据来源url
        self.pure_addr_list = list()

    # 获取数据来源url
    def get_pure_addr_list(self):
        # 获取页面源码
        source_html = self.get_response_content(self.url).decode('gbk')
        # 从源码中批量提取数据url
        all_addr = re.findall(r'wkbos\.bdimg\.com.*?json.*?expire.*?\}', source_html)
        if len(all_addr)==0:
            all_addr = re.findall(r'wkbjbos\.bdimg\.com.*?json.*?expire.*?\}', source_html)
        pure_addr_list = list()
        # 获取文档标题
        self.title = etree.HTML(source_html).xpath("//title/text()")[0].strip()
        # 净化数据来源url列表
        for addr in all_addr:
            addr = "https://" + addr.replace("\\\\\\/", "/")
            addr = addr[:-5]
            pure_addr_list.append(addr)
        self.pure_addr_list = pure_addr_list

        return pure_addr_list

    # 从数据来源的url列表中提取数据
    def get_json_content(self, url_list):
        result = ''
        flag=0
        temp_list=[]
        for pure_addr in url_list:
            try:

                # 获取json数据
                content = self.get_response_content(pure_addr).decode()
                # 处理json数据
                content = re.match(r'.*?\((.*)\)$', content).group(1)
                # 将json数据中需要的内容提取出来
                temp_list.append(content)
            except Exception as e:
                print 'error'
                return '0'
                pass
        #print len(temp_list)
        #print temp_list
        for index in range(len(temp_list)/2):
            all_body_info = json.loads(temp_list[index])["body"]
            # 遍历获取所有信息,并将信息拼接到一起
            y_tmp = 0
            #print all_body_info
            #print'2error'
            if all_body_info==None:
                #print'2error'
                continue
            for body_info in all_body_info:
                try:
                    if y_tmp == body_info["p"].get("y") :
                      result = result + body_info["c"].strip()
                    else:
                        result = result +'\r\n'+ body_info["c"].strip()
                        y_tmp = body_info["p"].get("y")
                except Exception as e:
                    #print(e)
                    
                    pass

        # 创建文件目录
        result2 = result.replace('\r\n','')
        print('%s'%("doc length: "),len(result2))
        if len(result2)==0:
            return '0'
        try:
            path =  os.sep +'data'+os.sep+'www'+os.sep+'static'+os.sep+ self.docType
            os.makedirs(path)
        except Exception as e:
            print(e)
            pass
        # 创建文件,保存信息
        try:
            file_name = os.sep +'data'+os.sep+'www'+os.sep+'static'+os.sep+ self.docType+os.sep + self.title + '.docx'
            file=docx.Document()
            file.styles['Normal'].font.name = u'宋体'
            file.add_paragraph(result)
            #file.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            file.save(file_name)
            return os.sep+'static'+os.sep+ self.docType+os.sep + self.title + '.docx'
            #with codecs.open(file_name, 'w','utf-8') as f:
                #f.write(result)
                #print("%s %s"%("已经保存为:".decode('utf8'), self.title + '.txt'))
                #f.close()
        except Exception as e:
            print(e)
            print'othererror'
            return '0'


# 创建获取ppt的类
class BDWKPPT(BaiduWK):
    def __init__(self, url):
        self.all_img_url = list()
        super(BDWKPPT,self).__init__(url)
    # 获取json数据保存文件
    def get_ppt_json_info(self):
        # 获取源文件
        ppt_source_html = self.get_response_content(self.url)
        # 解析源文件
        content = ppt_source_html.decode('gbk')
        # print("-->",len(content))
        # 测试
        with open("test.html", "w") as f:
            f.write(content)
        # 获取文档Id
        self.docId = re.findall(r"docId.*?(\w{24}?)\'\,", content)[0]
        # 拼接请求json的接口
        source_json_url = 'https://wenku.baidu.com/browse/getbcsurl?doc_id=%s&type=ppt&callback=zhaozhao' % self.docId
        # 获取字符串类型的json数据
        str_source_json = self.get_response_content(source_json_url).decode()
        # 处理字符串类型的json数据,使其成为标准格式
        try:
            pure_str_source_json = re.match(r'.*?\((.*?)\)', str_source_json).group(1)

        # 将字符串json转为可处理的正式json
            source_json = json.loads(pure_str_source_json)

        # 遍历字典中的数据类型list
            for j in source_json['list']:
            # 创建临时列表
                temp_num_url = list()
            # 将url和page拼接到列表中
                temp_num_url.append(j["zoom"])
                temp_num_url.append(j["page"])
            # 将列表信息添加到全局变量中
                self.all_img_url.append(temp_num_url)

        # 建立文件夹
            try:
                path =  os.sep +'data'+os.sep+'www'+os.sep+'static'+os.sep+ self.docType
                os.makedirs(path)
            except Exception as e:
                # print("---->>",e)
                pass
            with zipfile.ZipFile(os.sep +'data'+os.sep+'www'+os.sep+'static'+os.sep+ self.docType+"/%s.zip"%(self.title),mode='w') as zipf:
                for img_url in self.all_img_url:
                # print(img_url)
                    print("正在获取第%d页资源(剩余%d页)".decode('utf8') % (img_url[1], len(self.all_img_url) - img_url[1]))
                    data = self.get_response_content(img_url[0])
                    path = "./%s" % ( str(img_url[1]) + '.jpg')
                    with open(path, 'wb') as f:
                        f.write(data)
                    zipf.write(path)
                    os.remove(path)
                #os.rmdir(os.getcwd()+"\\ppt\\%s"%(self.title))
            os.remove('test.html')    
            print("写入完毕".decode('utf8'))
            return os.sep+'static'+os.sep+ self.docType+"/%s.zip"%(self.title)
        except Exception as e:
                # print("---->>",e)
            print'不支持该类文档'    
            return '0'
            


def main1(url1):
    try:
        url = url1
        docType = BaiduWK(url).docType
    except:
        return [-1,'0','0',0]
        os._exit(0)

    print("%s %s"%("type is-->", docType))

    if docType == "ppt":

        ppt = BDWKPPT(url)
        print("%s %s"%("您将要获取的演示文稿(ppt)名称为:".decode('utf-8'), ppt.title))
        path = ppt.get_ppt_json_info()
        if (path!='0'):
            path2 = 'https://quezz.cn'+path
            return [1,path2,ppt.title,os.path.getsize(os.sep +'data'+os.sep+'www'+path)]
        else:
            return [0,'0','0',0]
    elif docType == "doc" or docType== 'pdf':
        word = BDWKDOC(url)
        print("%s %s"%("您将要获取的word名称为:".decode('utf8'), word.title))
        pure_addr_list = word.get_pure_addr_list()
        print len(pure_addr_list)
        path = word.get_json_content(pure_addr_list)
        if (path!='0'):
            path2 = 'https://quezz.cn'+path
            return [2,path2,word.title,os.path.getsize(os.sep +'data'+os.sep+'www'+path)]
        else:
            return [0,'0','0',0]
    # elif docType == "pdf":
    #     pdf = BDWKDOC(url)
    #     print("%s %s" % ("您将要获取的pdf名称为:".decode('utf8'),  pdf.title))
    #     pure_addr_list = pdf.get_pure_addr_list()
    #     pdf.get_json_content(pure_addr_list)

    elif docType == "txt":

        txt = BDWKTXT(url)
        print("%s %s"% ("您将要下载的文本文档(txt)名称为:".decode('utf8'), txt.title))
        path = txt.get_txt(url)
        print path
        if (path!='0'):
            path2 = 'https://quezz.cn'+path
            return [3,path2,txt.title,os.path.getsize(os.sep +'data'+os.sep+'www'+path)]
        else:
            return [0,'0','0',0]